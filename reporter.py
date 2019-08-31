# 把屏幕输出文本、保存的图片写入到可作为报告的word文档中
from docx import Document
from docx.shared import Inches, Pt

doc = Document()  # 以默认模板建立文档对象

doc.add_heading('Level report', 0)

'''
def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size
'''

distance = Inches(2)
sec = doc.sections[0]  # sections对应文档中的“节”
sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距
sec.right_margin = distance
sec.top_margin = distance
sec.bottom_margin = distance
sec.page_width = Inches(12)  # 设置页面宽度
sec.page_height = Inches(20)  # 设置页面高度


#添加段落文本
fileobj = open('level_report.txt', 'r') # 读取txt文本的内容，存到字符串
try:
    paragraphstr = fileobj.read()
finally:
    fileobj.close()

run = doc.add_paragraph().add_run(paragraphstr)
#paragraph = doc.add_paragraph(paragraphstr,run.font.size)
run.font.name = 'Calibri'
run.font.size = Pt(20)

#ph_format = paragraph.paragraph_format
#ph_format.space_before = Pt(10)  # 设置段前间距
#ph_format.space_after = Pt(12)  # 设置段后间距
#ph_format.line_spacing = Pt(19)  # 设置行间距

#添加图片
doc.add_picture('IL of pipe lagging.png', width=Inches(8.0))

#保存word文档
doc.save('level_report.docx')